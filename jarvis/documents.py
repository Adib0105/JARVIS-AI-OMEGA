from __future__ import annotations

import csv
import hashlib
import io
import re
from pathlib import Path

from .local_files import LocalFiles
from .security.secrets import redact_secrets


DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.xlsm', '.csv', '.txt', '.md'}


class DocumentReader:
    def __init__(self, files: LocalFiles | None = None):
        self.files = files or LocalFiles()

    def _safe_path(self, file_path: str) -> Path:
        path = Path(file_path).expanduser().resolve()
        if not self.files._is_inside_root(path):
            raise PermissionError('Document is outside approved roots.')
        if self.files._looks_secret(path):
            raise PermissionError('Secret-like path is blocked.')
        if not path.is_file():
            raise FileNotFoundError(path)
        if path.suffix.lower() not in DOCUMENT_EXTENSIONS:
            raise PermissionError(f'Unsupported document type: {path.suffix}')
        return path

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open('rb') as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b''):
                digest.update(chunk)
        return digest.hexdigest()

    def extract(self, file_path: str, max_chars: int = 120000) -> dict:
        path = self._safe_path(file_path)
        cap = max(2000, min(int(max_chars), 250000))
        suffix = path.suffix.lower()
        stat = path.stat()
        content_sha256 = self._sha256(path)

        if suffix == '.pdf':
            text, meta = self._pdf(path, cap)
        elif suffix == '.docx':
            text, meta = self._docx(path, cap)
        elif suffix in {'.xlsx', '.xlsm'}:
            text, meta = self._xlsx(path, cap)
        elif suffix == '.csv':
            text, meta = self._csv(path, cap)
        else:
            text = path.read_text(encoding='utf-8', errors='replace')[:cap]
            meta = {'type': suffix.lstrip('.'), 'characters': len(text)}

        raw_extracted = text[:cap]
        extracted, secret_findings = redact_secrets(raw_extracted)
        extracted_sha256 = hashlib.sha256(extracted.encode('utf-8', errors='replace')).hexdigest()
        meta = dict(meta)
        meta['characters'] = len(extracted)
        if secret_findings:
            meta['credential_redactions'] = len(secret_findings)
            meta['credential_redaction_types'] = sorted({item.description for item in secret_findings})

        meaningful = self._meaningful_text(extracted)
        if suffix == '.pdf' and len(meaningful) < 80:
            meta['limited_extraction'] = True
            meta['warning'] = (
                'Very little selectable text was found in this PDF. It may be scanned/image-based; '
                'JARVIS learned only the text that could be extracted safely.'
            )
        else:
            meta['limited_extraction'] = False

        meta.update({
            'content_sha256': content_sha256,
            'extracted_sha256': extracted_sha256,
            'mtime_ns': int(stat.st_mtime_ns),
        })
        return {
            'path': str(path),
            'name': path.name,
            'size_bytes': stat.st_size,
            'mtime_ns': int(stat.st_mtime_ns),
            'content_sha256': content_sha256,
            'extracted_sha256': extracted_sha256,
            'metadata': meta,
            'text': extracted,
        }

    @staticmethod
    def _meaningful_text(text: str) -> str:
        # Ignore synthetic PDF page markers when judging whether a document really
        # contained extractable text. A one-page scanned PDF otherwise appears to
        # have ~16 characters only because of "--- PAGE 1 ---".
        value = re.sub(r'---\s*PAGE\s+\d+\s*---', ' ', str(text or ''), flags=re.IGNORECASE)
        return re.sub(r'\s+', ' ', value).strip()

    @staticmethod
    def _pdf(path: Path, cap: int) -> tuple[str, dict]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts: list[str] = []
        total = 0
        for idx, page in enumerate(reader.pages):
            chunk = page.extract_text() or ''
            header = f'\n--- PAGE {idx + 1} ---\n'
            parts.append(header + chunk)
            total += len(header) + len(chunk)
            if total >= cap:
                break
        text = ''.join(parts)[:cap]
        return text, {'type': 'pdf', 'pages': len(reader.pages), 'characters': len(text)}

    @staticmethod
    def _docx(path: Path, cap: int) -> tuple[str, dict]:
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
            if sum(map(len, parts)) >= cap:
                break
        text = '\n'.join(parts)[:cap]
        return text, {'type': 'docx', 'paragraphs': len(doc.paragraphs), 'characters': len(text)}

    @staticmethod
    def _xlsx(path: Path, cap: int) -> tuple[str, dict]:
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            out = io.StringIO()
            row_count = 0
            for ws in wb.worksheets[:20]:
                out.write(f'\n--- SHEET: {ws.title} ---\n')
                for row in ws.iter_rows(values_only=True):
                    values = [str(v) if v is not None else '' for v in row]
                    out.write('\t'.join(values).rstrip() + '\n')
                    row_count += 1
                    if out.tell() >= cap or row_count >= 10000:
                        break
                if out.tell() >= cap or row_count >= 10000:
                    break
            text = out.getvalue()[:cap]
            return text, {'type': 'xlsx', 'sheets': len(wb.sheetnames), 'rows_read': row_count, 'characters': len(text)}
        finally:
            wb.close()

    @staticmethod
    def _csv(path: Path, cap: int) -> tuple[str, dict]:
        out = io.StringIO()
        rows = 0
        with path.open('r', encoding='utf-8-sig', errors='replace', newline='') as handle:
            reader = csv.reader(handle)
            for row in reader:
                out.write('\t'.join(row) + '\n')
                rows += 1
                if out.tell() >= cap or rows >= 20000:
                    break
        text = out.getvalue()[:cap]
        return text, {'type': 'csv', 'rows_read': rows, 'characters': len(text)}
